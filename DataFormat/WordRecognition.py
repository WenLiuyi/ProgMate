"""
# Word文件读取
Author: shq
Packages: os
"""
import os, win32com.client
from docx import Document

def all_words(file_dir):
    """
    # 获取文件夹下所有word路径
    """
    word_list=[]
    for root, dirs, files in os.walk(file_dir):
        for file in files:
            if file.endswith('.docx') or file.endswith('.doc'): # 检验文件拓展名
                file_path = os.path.join(root, file)
                word_list.append(file_path)
    return word_list

def read_word(word_dir):
    """
    # 读取word文件的文本
    """
    words = all_words(word_dir)
    result_texts = []
    for word in words:
        if word.split('.')[-1] == 'docx': # 新版word
            doc = Document(word)
            document_text = ""
            for paragraph in doc.paragraphs:
                document_text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        document_text += cell.text + "\t"
                    document_text += "\n"
        if word.split('.')[-1] == 'doc': #旧版word
            word = win32com.client.Dispatch("Word.Application")
            doc = word.Documents.Open(word)
            for para in doc.Paragraphs:
                document_text += paragraph.text + "\n"
                print(para.Range.Text)
            word.Quit()
        result_texts.append(document_text)
        if not os.path.exists('./text/'):
            os.makedirs('./text/')
        with open('./text/'+os.path.basename(word)+'_text.txt', 'w', encoding='utf-8') as file:
            file.write(document_text)
    return result_texts
 
if __name__ == "__main__":
    word_dir = ".\\testfiles"
    texts = read_word(word_dir)
    for i, text in enumerate(texts, start=1):
        print(f"Document {i} Content:\n{text}\n")
